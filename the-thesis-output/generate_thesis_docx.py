#!/usr/bin/env python3
"""
Generate complete thesis DOCX from VitePress markdown sources.
Assembles all chapters, appendices, references, figures, and front matter.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
DOCS = os.path.join(os.path.dirname(BASE), "docs")
THESIS = os.path.join(DOCS, "thesis")
APPENDICES = os.path.join(DOCS, "appendices")
FIGURES = os.path.join(DOCS, "figures")
OUTPUT = os.path.join(BASE, "The Thesis.docx")


def read_md(filename):
    path = os.path.join(THESIS, filename)
    if not os.path.exists(path):
        path = os.path.join(APPENDICES, filename)
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    # Strip YAML frontmatter
    content = re.sub(r"^---\n.*?\n---\n*", "", content, flags=re.DOTALL)
    return content.strip()


def setup_styles(doc):
    """Configure document styles for academic thesis formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for level, size, bold in [
        ("Heading 1", 16, True),
        ("Heading 2", 14, True),
        ("Heading 3", 13, True),
        ("Heading 4", 12, True),
        ("Heading 5", 12, True),
    ]:
        s = doc.styles[level]
        s.font.name = "Times New Roman"
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(18 if level in ("Heading 1", "Heading 2") else 12)
        s.paragraph_format.space_after = Pt(6)

    # Caption style
    if "Caption" not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["Caption"]
    cap.font.name = "Times New Roman"
    cap.font.size = Pt(10)
    cap.font.italic = True
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)


def add_title_page(doc):
    """Create the thesis title page."""
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("POLITECNICO DI MILANO")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("School of Management Engineering")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Master of Science in Management Engineering")
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Decentralised Applications in Focus:\n"
        "Governance, Market Structure, and Adoption Patterns"
    )
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Thesis presented for the degree of\nMaster of Science in Management Engineering")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    for _ in range(2):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Written by\nKristian Kremenov Kirilov\nStudent ID: 10855006")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Academic Year: 2024–2025")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_page_break()


def add_toc_placeholder(doc):
    """Add a Table of Contents placeholder page."""
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph(
        "[Table of Contents will be generated automatically in Word. "
        "Right-click here and select 'Update Field' after opening in Word.]"
    )
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc.add_page_break()


def add_list_of_figures(doc):
    """Add List of Figures placeholder."""
    doc.add_heading("List of Figures", level=1)
    figures = [
        ("Figure 3.1", "Data collection and linkage pipeline"),
        ("Figure 3.2", "Sample construction funnel"),
        ("Figure 4.1", "Governance label distribution — loose universe (N=834)"),
        ("Figure 4.2", "Governance × ownership heatmap — loose backtest universe (N=834)"),
        ("Figure 4.3", "Governance × token type heatmap — loose backtest universe (N=834)"),
        ("Figure 4.4", "Governance label distribution — strict universe (N=68)"),
        ("Figure 4.5", "Governance × ownership heatmap — strict sample (N=68)"),
        ("Figure 4.6", "Governance type × token type heatmap — strict sample (N=68)"),
        ("Figure 4.7", "Market capitalisation and user concentration — strict sample (N=68)"),
        ("Figure 4.8", "Market dynamics — loose universe comparison (N=834)"),
        ("Figure 4.9", "Top-15 blockchain deployments — strict sample (N=68)"),
        ("Figure 4.11", "Sector-level performance metrics — strict sample (N=68)"),
        ("Figure 4.13", "Governance × ownership heatmap — sector × category cohort"),
        ("Figure 4.14", "Governance × token type heatmap — sector × category cohort"),
        ("Figure 4.15", "Governance label distribution — sector × category cohort"),
        ("Figure 5.1", "Governance type distribution (strict sample, N=68)"),
        ("Figure 5.2", "Market cap and user base concentration (strict sample)"),
        ("Figure 5.3", "Engagement vs. economic value by vertical"),
        ("Figure 5.4", "Multi-chain deployment and performance pathway"),
    ]
    for num, title in figures:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}: ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(title).font.size = Pt(11)
    doc.add_page_break()


def add_list_of_tables(doc):
    """Add List of Tables placeholder."""
    doc.add_heading("List of Tables", level=1)
    tables = [
        ("Table 3.1", "Sample construction at a glance"),
        ("Table 3.2", "Target Ecosystem Mapping"),
        ("Table 3.3", "Variable categories and counts"),
        ("Table 3.4", "Governance score component weights"),
        ("Table 3.5", "Sample sizes in governance-depth studies"),
        ("Table 4.1", "Headline metrics comparison: loose universe versus strict high-signal sample"),
        ("Table 4.2", "Cross-tabulation: decentralisation level × governance type, strict sample (N=68)"),
        ("Table 4.4", "Top-10 chains by DeFi DApp presence (loose universe, N=105)"),
        ("Table 4.5", "Top DeFi DApps by active users (loose universe)"),
        ("Table 4.6", "DeFi governance distribution (loose universe, N=105)"),
        ("Table 4.7", "Top-5 chains by prediction market DApp presence"),
        ("Table 4.8", "Top prediction market DApps by active users"),
        ("Table 4.9", "Prediction market governance distribution"),
        ("Table 4.10", "AI DApp distribution by application category"),
        ("Table 4.11", "Top chains by AI DApp deployment"),
        ("Table 4.12", "Representative AI DApps by active users"),
        ("Table 4.13", "AI DApp governance distribution"),
        ("Table 4.14", "Decentralisation levels: AI DApps versus full ecosystem"),
        ("Table 4.15", "AI DApp token type distribution"),
        ("Table 4.19", "Summary of anomaly categories, prevalence, and cross-references"),
        ("Table 4.20", "Cross-ecosystem structural comparison (loose universe)"),
        ("Table 4.21", "Sector composition of the strict sample (N=68)"),
    ]
    for num, title in tables:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}: ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(title).font.size = Pt(11)
    doc.add_page_break()


def add_list_of_symbols(doc):
    """Add List of Symbols and Abbreviations."""
    doc.add_heading("List of Symbols and Abbreviations", level=1)
    symbols = [
        ("AMM", "Automated Market Maker"),
        ("API", "Application Programming Interface"),
        ("BIS", "Bank for International Settlements"),
        ("BNB", "Build and Build (Binance Smart Chain)"),
        ("CMC", "CoinMarketCap"),
        ("DAO", "Decentralised Autonomous Organisation"),
        ("DApp", "Decentralised Application"),
        ("DeFi", "Decentralised Finance"),
        ("DePIN", "Decentralised Physical Infrastructure Network"),
        ("DEX", "Decentralised Exchange"),
        ("EIP", "Ethereum Improvement Proposal"),
        ("EOA", "Externally Owned Account"),
        ("ERC", "Ethereum Request for Comments"),
        ("EVM", "Ethereum Virtual Machine"),
        ("FDV", "Fully Diluted Valuation"),
        ("GameFi", "Game Finance"),
        ("HHI", "Herfindahl–Hirschman Index"),
        ("ICO", "Initial Coin Offering"),
        ("IPFS", "InterPlanetary File System"),
        ("L1", "Layer 1 (base blockchain)"),
        ("L2", "Layer 2 (scaling solution)"),
        ("NFT", "Non-Fungible Token"),
        ("P2E", "Play-to-Earn"),
        ("PCA", "Principal Component Analysis"),
        ("PoS", "Proof of Stake"),
        ("PoW", "Proof of Work"),
        ("RPC", "Remote Procedure Call"),
        ("RWA", "Real World Assets"),
        ("SoK", "Systematization of Knowledge"),
        ("TVL", "Total Value Locked"),
        ("UAW", "Unique Active Wallets"),
        ("UIP", "Uniswap Improvement Proposal"),
        ("N", "Sample size"),
        ("α", "Power-law exponent"),
        ("κ", "Cohen’s kappa (inter-rater reliability)"),
        ("ρ", "Spearman’s rank correlation coefficient"),
        ("r", "Correlation coefficient"),
        ("k", "Number of clusters (K-means)"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Symbol / Abbreviation", "Definition"]):
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
    for sym, defn in symbols:
        row = table.add_row().cells
        row[0].text = sym
        row[1].text = defn
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
    doc.add_page_break()


def parse_table(lines):
    """Parse a markdown table into list of rows (each row is list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        # Skip separator rows
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    """Add a parsed markdown table to the document."""
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.add_row()
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = row.cells[j]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Times New Roman"
                        if i == 0:
                            run.bold = True


def try_add_image(doc, img_ref):
    """Try to add an image from a markdown image reference."""
    # Extract path from markdown image syntax
    match = re.search(r"\((.*?)\)", img_ref)
    if not match:
        return False
    rel_path = match.group(1)
    # Resolve relative to docs/
    if rel_path.startswith("../"):
        img_path = os.path.join(THESIS, rel_path)
    elif rel_path.startswith("/"):
        img_path = os.path.join(DOCS, rel_path.lstrip("/"))
    else:
        img_path = os.path.join(DOCS, rel_path)

    img_path = os.path.normpath(img_path)
    if os.path.exists(img_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            return True
        except Exception:
            pass
    return False


def format_inline(paragraph, text):
    """Add text with basic inline formatting (bold, italic) to a paragraph."""
    # Process bold and italic patterns
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def process_markdown(doc, content, base_heading_level=1):
    """Process markdown content and add to document."""
    lines = content.split("\n")
    i = 0
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code blocks (mermaid diagrams -> skip, other code -> format)
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                # Add code block content
                if code_lines and not any("mermaid" in cl for cl in [code_lines[0]] if code_lines):
                    for cl in code_lines:
                        p = doc.add_paragraph()
                        run = p.add_run(cl)
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                code_lines = []
                i += 1
                continue
            else:
                lang = line.strip()[3:]
                if lang == "mermaid":
                    # Skip mermaid blocks - add placeholder (don't set in_code_block)
                    i += 1  # move past opening ```mermaid
                    while i < len(lines) and lines[i].strip() != "```":
                        i += 1
                    if i < len(lines):
                        i += 1  # skip closing ```
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run("[Diagram — see VitePress version for interactive rendering]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    continue
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            rows = parse_table(table_lines)
            add_table_to_doc(doc, rows)
            doc.add_paragraph("")  # spacing after table
            in_table = False
            table_lines = []
            # Don't increment i, process current line

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,5})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Clean markdown from heading
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            text = re.sub(r"\*(.*?)\*", r"\1", text)
            # Map to document heading levels
            doc_level = min(level, 4)
            doc.add_heading(text, level=doc_level)
            i += 1
            continue

        # Images
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img_match:
            alt_text = img_match.group(1)
            try_add_image(doc, line.strip())
            i += 1
            continue

        # Figure captions (italic lines starting with Figure or *)
        if line.strip().startswith("*Figure") or line.strip().startswith("*Fig"):
            text = line.strip().strip("*").strip()
            p = doc.add_paragraph(style="Caption")
            p.add_run(text).font.italic = True
            i += 1
            continue

        # Horizontal rules
        if re.match(r"^---+\s*$", line.strip()):
            i += 1
            continue

        # Bullet points
        bullet_match = re.match(r"^(\s*)([-*])\s+(.*)", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(3)
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 3:
                p.paragraph_format.left_indent = Inches(0.5)
            format_inline(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r"^(\s*)(\d+)\.\s+(.*)", line)
        if num_match:
            text = num_match.group(3)
            p = doc.add_paragraph(style="List Number")
            format_inline(p, text)
            i += 1
            continue

        # Block quotes
        if line.strip().startswith(">"):
            text = line.strip().lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            format_inline(p, text)
            p.runs[0].font.italic = True if p.runs else None
            i += 1
            continue

        # Math formulas ($$..$$)
        if line.strip().startswith("$$"):
            formula = line.strip().strip("$").strip()
            if not formula:
                # Multi-line formula
                i += 1
                formula_parts = []
                while i < len(lines) and not lines[i].strip().startswith("$$"):
                    formula_parts.append(lines[i].strip())
                    i += 1
                formula = " ".join(formula_parts)
                i += 1  # skip closing $$
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            i += 1
            continue

        # Regular paragraph - collect continuation lines
        para_text = line.strip()
        while (
            i + 1 < len(lines)
            and lines[i + 1].strip()
            and not lines[i + 1].strip().startswith("#")
            and not lines[i + 1].strip().startswith("|")
            and not lines[i + 1].strip().startswith("!")
            and not lines[i + 1].strip().startswith("```")
            and not lines[i + 1].strip().startswith("- ")
            and not lines[i + 1].strip().startswith("* ")
            and not re.match(r"^\d+\.\s", lines[i + 1].strip())
            and not lines[i + 1].strip().startswith(">")
            and not lines[i + 1].strip().startswith("$$")
            and not re.match(r"^---+\s*$", lines[i + 1].strip())
            and not lines[i + 1].strip().startswith("*Figure")
            and not lines[i + 1].strip().startswith("*Fig")
        ):
            i += 1
            para_text += " " + lines[i].strip()

        if para_text:
            p = doc.add_paragraph()
            format_inline(p, para_text)

        i += 1

    # Flush remaining table
    if in_table and table_lines:
        rows = parse_table(table_lines)
        add_table_to_doc(doc, rows)


def build_bibliography(doc):
    """Build the full bibliography from all chapter references."""
    doc.add_heading("Bibliography", level=1)

    # Collect all references from all chapters
    refs = set()
    for md_file in [
        "01-introduction.md",
        "02-literature-review.md",
        "03-methodology.md",
        "04-results.md",
        "04c-case-studies.md",
        "05-discussion.md",
        "06-conclusions.md",
    ]:
        path = os.path.join(THESIS, md_file)
        if not os.path.exists(path):
            continue
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        # Find reference sections
        ref_match = re.search(
            r"(?:^##+ References.*?\n)(.*?)(?=\n---|\n##[^#]|\Z)",
            content,
            re.MULTILINE | re.DOTALL,
        )
        if ref_match:
            ref_block = ref_match.group(1)
            # Split by blank lines to get individual references
            for ref in re.split(r"\n\n+", ref_block):
                ref = ref.strip()
                if ref and not ref.startswith("#") and not ref.startswith("---"):
                    # Clean markdown
                    ref = re.sub(r"\*([^*]+)\*", r"\1", ref)
                    refs.add(ref)

    # Sort alphabetically
    sorted_refs = sorted(refs, key=lambda x: x.lower())

    for ref in sorted_refs:
        if ref.strip():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            format_inline(p, ref)
            for run in p.runs:
                run.font.size = Pt(11)


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)  # A4
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    setup_styles(doc)

    # ==================== FRONT MATTER ====================
    add_title_page(doc)
    add_toc_placeholder(doc)

    # ==================== ABSTRACT ====================
    abstract_md = read_md("abstract.md")
    process_markdown(doc, abstract_md)
    doc.add_page_break()

    # ==================== LISTS ====================
    add_list_of_figures(doc)
    add_list_of_tables(doc)
    add_list_of_symbols(doc)

    # ==================== CHAPTERS ====================
    chapters = [
        ("01-introduction.md", "Chapter 1 — Introduction"),
        ("02-literature-review.md", "Chapter 2 — Literature Review"),
        ("03-methodology.md", "Chapter 3 — Methodology"),
        ("04-results.md", "Chapter 4 — Results"),
        ("04c-case-studies.md", None),  # Continuation of Chapter 4
        ("05-discussion.md", "Chapter 5 — Discussion"),
        ("06-conclusions.md", "Chapter 6 — Conclusions"),
    ]

    for filename, chapter_title in chapters:
        content = read_md(filename)
        # Remove the chapter's own reference section for the main body
        # (references go in the bibliography)
        content = re.sub(
            r"\n### References.*$",
            "",
            content,
            flags=re.DOTALL,
        )
        content = re.sub(
            r"\n## References for Chapter.*$",
            "",
            content,
            flags=re.DOTALL,
        )
        content = re.sub(
            r"\n## References for .*$",
            "",
            content,
            flags=re.DOTALL,
        )
        # Remove word count lines
        content = re.sub(r"\n\*Word count:.*?\*", "", content)
        content = re.sub(r"\n\*Status:.*?\*", "", content)

        process_markdown(doc, content)
        doc.add_page_break()

    # ==================== BIBLIOGRAPHY ====================
    build_bibliography(doc)
    doc.add_page_break()

    # ==================== APPENDICES ====================
    doc.add_heading("Appendices", level=1)
    doc.add_paragraph("")

    # Appendix A - Variable Codebook
    codebook_md = read_md("variable-codebook.md")
    process_markdown(doc, codebook_md)
    doc.add_page_break()

    # Appendix B - Analytical Pipeline
    pipeline_md = read_md("analytical-pipeline.md")
    process_markdown(doc, pipeline_md)

    # ==================== SAVE ====================
    doc.save(OUTPUT)
    print(f"Thesis saved to: {OUTPUT}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
